import datetime
import os

import docx
from loguru import logger

STOP_SYMBOL = ['/', ' ', '*', '+', '=', '-', ':', ';', '"', '%']
CALC_SYMBOL = ['/', '*', '+', '-', '**', '%']
name_temp = 'temp.docx'



@logger.catch
def read_doc():
    # получение списка уникальных параметров
    doc = docx.Document(name_temp)

    len_parag = len(doc.paragraphs)
    param_name_list = []     # список переменных для ввода
    param_name_list_all = [] # общий список переменных
    for i in range(len_parag):
        param_name = ''
        len_text = len(doc.paragraphs[i].text)
        param_enter = False
        # поиск символов в строке
        for ii in range(len_text):
            read_symbol = doc.paragraphs[i].text[ii]

            if read_symbol in STOP_SYMBOL:
                if read_symbol == '=':
                    param_enter = True

                if param_name != '':
                    if not param_enter:
                        param_name_list.append(param_name)
                    param_name_list_all.append(param_name)
                    param_name = ''
            else:
                param_name += read_symbol
        if param_name != '':
            param_name_list_all.append(param_name)

    param_name_list_unique_all = set(param_name_list_all)
    param_name_list_unique = set(param_name_list)
    param_name_list_unique = sorted(param_name_list_unique_all-param_name_list_unique)
    param_name_list_unique_all = sorted(param_name_list_unique_all)

    list_all = []
    list_all.append(param_name_list_unique)
    list_all.append(param_name_list_unique_all)

    return list_all



@logger.catch
async def write_doc(dict_param, id_calc, calc_date_create):
    doc_read = docx.Document(name_temp)

    doc_write = docx.Document()
    doc_write.add_heading('Файл некого расчета!',)
    doc_write.add_paragraph('Формулы расчета:')

    for i in range(len(doc_read.paragraphs)):
        doc_write.add_paragraph(doc_read.paragraphs[i].text)
    doc_write.add_paragraph('')
    doc_write.add_paragraph('Введенные значения:')


    for key in dict_param:
        if dict_param[key] != None:
            doc_write.add_paragraph(f'{key}: {dict_param[key]}')
    doc_write.add_paragraph('')
    doc_write.add_paragraph('Полученные значения:')

    for i in range(len(doc_read.paragraphs)):
        param_name = ''
        str_paragraphs = ''
        key_calc = ''
        for ii in range(len(doc_read.paragraphs[i].text)):
            read_symbol = doc_read.paragraphs[i].text[ii]
            if read_symbol in STOP_SYMBOL:
                if read_symbol in CALC_SYMBOL:
                     str_paragraphs += read_symbol
                else:
                    if param_name != '':
                        if key_calc == '':
                            key_calc = param_name
                        else:
                            str_paragraphs += f' {dict_param[param_name]} '
                        param_name = ''
            else:
                param_name += read_symbol
        if param_name != '':
            str_paragraphs += f' {dict_param[param_name]} '
        try:
            calc_value = eval(str_paragraphs)
        except:
            print('ошибка преобразования расчета')
        dict_param[key_calc] = calc_value

        doc_write.add_paragraph(f'{key_calc}: {dict_param[key_calc]}')


    dir_save = f'{dir_temp}{id_calc}/'
    os.mkdir(dir_save)
    name_file_calc = f'{dir_save}Расчет от {calc_date_create}.docx'
    doc_write.save(name_file_calc)
    return dir_save, name_file_calc


if __name__ == '__main__':
    x = read_doc()
    print(x)
    dir_temp = ''
else:
    dir_temp = 'word/'

name_temp = dir_temp+name_temp